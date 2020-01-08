from typing import List, Tuple

from docx import Document

import consts
from file_utils import safe_join_path
from get_times_from_yeshiva_site import get_times_as_titles_and_times, convert_month_to_month_number, is_year_leaped, \
    get_shabat_times, convert_to_time_of_day, get_start_of_months
from time_of_day import TimeOfDay, Time


def get_specific_day_times(day_to_extract: str, time_of_days: List[TimeOfDay]) -> List[TimeOfDay]:
    return [day_times for day_times in time_of_days if day_times[consts.DAY_IN_WEEK] == day_to_extract]


def get_last_month(month_number: int, year_number: int) -> Tuple[int, int]:
    if month_number == min(consts.LEAP_YEAR_MONTH_NUMBERS.values()):
        if is_year_leaped(year_number):
            return max(consts.LEAP_YEAR_MONTH_NUMBERS.values()), year_number - 1
        return max(consts.NON_LEAP_YEAR_MONTH_NUMBERS.values()), year_number - 1
    return month_number - 1, year_number


def get_next_month(month_number: int, year_number: int) -> Tuple[int, int]:
    if is_year_leaped(year_number):
        if month_number == len(consts.LEAP_YEAR_MONTH_NUMBERS):
            return min(consts.LEAP_YEAR_MONTH_NUMBERS.values()), year_number + 1
        return month_number + 1, year_number
    if month_number == len(consts.NON_LEAP_YEAR_MONTH_NUMBERS):
        return min(consts.NON_LEAP_YEAR_MONTH_NUMBERS.values()), year_number + 1
    return month_number + 1, year_number


def reverse_dict(dict_to_reverse: dict) -> dict:
    return {value: key for key, value in dict_to_reverse.items()}


def calculate_month_name(month_number: int, year_number: int) -> str:
    if is_year_leaped(year_number):
        return reverse_dict(consts.LEAP_YEAR_MONTH_NUMBERS)[month_number]
    return reverse_dict(consts.NON_LEAP_YEAR_MONTH_NUMBERS)[month_number]


def put_quotes(un_quoted: str) -> str:
    if len(un_quoted) == 1:
        return un_quoted + "'"
    return un_quoted[:-1] + '"' + un_quoted[-1]


def get_date_string(day_in_month: str, month: str) -> str:
    return put_quotes(day_in_month) + ' ב' + month


def is_shabat_in_month(shabat_times: TimeOfDay, month: str) -> bool:
    shabat_date = shabat_times[consts.DATE]
    if month in ['חשון', 'סיון']:
        month = month.replace('ו', 'וו')
    if month == 'כסלו':
        month = 'כסליו'
    return consts.SHABAT in shabat_date and month in shabat_date


def is_moladot_time_in_month(moladot_time: TimeOfDay, month: str) -> bool:
    return month == moladot_time[consts.HEBREW_MONTH]


def convert_plag_to_minha(plag_time: Time) -> Time:
    return plag_time - 15 - plag_time % 15


def convert_shkia_to_come_shabat(shkia_time: Time) -> Time:
    return shkia_time - 30 - shkia_time % 15


def get_min_time(day_times: List[TimeOfDay], time_name: str) -> Time:
    return min(day_time[time_name] for day_time in day_times)


def get_max_time(day_times: List[TimeOfDay], time_name: str) -> Time:
    return max(day_time[time_name] for day_time in day_times)


def calculate_minha_time_according_to_plag(friday_times: List[TimeOfDay]) -> Time:
    plag_min_time = get_min_time(friday_times, consts.PLAG)
    return convert_plag_to_minha(plag_min_time)


def calculate_come_shabat_time_according_to_shkia(friday_times: List[TimeOfDay]) -> Time:
    sun_set_min_time = get_min_time(friday_times, consts.SUN_SET)
    return convert_shkia_to_come_shabat(sun_set_min_time)


def convert_to_good_date_format(date_format: str) -> str:
    words = date_format.split()
    words[1] = words[1].replace(',', '')
    return words[0] + ' ' + words[1] + ' ' + words[2] + words[3] + ' ' + words[4]


def calculate_moladot_of_month(month_moladot: TimeOfDay) -> str:
    lines = month_moladot[consts.MOLADOT_TIME].splitlines(keepends=False)
    lines[consts.GOOD_DATE_FORMAT_LINE] = convert_to_good_date_format(lines[consts.GOOD_DATE_FORMAT_LINE])
    del lines[consts.BAD_DATE_FORMAT_LINE]
    return ' '.join(lines)


def is_according_to_plag(friday_times: List[TimeOfDay]) -> bool:
    plag_min_time = get_min_time(friday_times, consts.PLAG)
    minha_according_to_plag = convert_plag_to_minha(plag_min_time)
    return minha_according_to_plag.hours >= consts.HOUR_TO_GO_ACCORDING_TO_SHKIA


def get_fathers_and_sons_time(shabat_times: List[TimeOfDay]) -> Time:
    sun_set_min_time = get_min_time(shabat_times, consts.SUN_SET)
    sun_set_max_time = get_max_time(shabat_times, consts.SUN_SET)
    shabat_minha_min_time = sun_set_min_time - consts.SHABAT_MINHA_DURATION
    shabat_minha_max_time = sun_set_max_time - consts.SHABAT_MINHA_DURATION
    fathers_and_sons_min_time = (
            shabat_minha_max_time - consts.NOON_LESSON_MAX_DURATION - consts.FATHERS_AND_SONS_DURATION
    )
    fathers_and_sons_max_time = (
            shabat_minha_min_time - consts.NOON_LESSON_MIN_DURATION - consts.FATHERS_AND_SONS_DURATION
    )
    fathers_and_sons_min_time += 15 - fathers_and_sons_min_time % 15
    fathers_and_sons_max_time -= fathers_and_sons_max_time % 15

    if fathers_and_sons_min_time > fathers_and_sons_max_time:
        raise Exception(
            f'fathers and sons min time ({fathers_and_sons_min_time.strftime(consts.TIME_FORMAT)}) '
            f'was bigger than the max time ({fathers_and_sons_max_time.strftime(consts.TIME_FORMAT)})'
        )

    return max(fathers_and_sons_max_time - ((shabat_minha_min_time - '15:45') // 30 * 15), fathers_and_sons_min_time)


def convert_by_gimatria(word: str) -> int:
    return sum(consts.GIMATRIA_DICT.get(c, 0) for c in word)


def get_year_from_number(year_number: int) -> str:
    year = ''
    for gimatria_of_letter in sorted(consts.REVERSE_GIMATRIAA_DICT, reverse=True):
        while year_number >= gimatria_of_letter:
            year += consts.REVERSE_GIMATRIAA_DICT[gimatria_of_letter]
            year_number -= gimatria_of_letter
    return year


def get_fields_titles_according_to_plag():
    yield 'פרשת השבוע'
    yield 'תאריך'
    yield 'תאריך לועזי'
    yield 'מנחה של שישי'
    yield 'פלג המנחה'
    yield 'בואי כלה'
    yield 'הדלקת נרות'
    yield 'שקיעה'
    yield 'שיעור בוקר שבת'
    yield 'שחרית של שבת'
    yield 'סוף זמן ק"ש'
    yield 'אבות ובנים'
    yield 'שיעור אחה"צ שבת'
    yield 'מנחה של שבת'
    yield 'צאת שבת רש"י (גאונים)'
    yield 'צאת שבת ר"ת'


def get_fields_to_write_according_to_plag(
        friday_time, shabat_time, shabat_special_time, month, friday_times, shabat_times
):
    friday_minha_time = calculate_minha_time_according_to_plag(friday_times)
    fathers_and_sons_time = get_fathers_and_sons_time(shabat_times)
    # פרשות
    yield shabat_special_time[consts.PARASHA]
    # תאריך עברי
    yield get_date_string(shabat_time[consts.DAY_IN_MONTH], month)
    # תאריך לועזי
    yield shabat_time[consts.WEIRD_DATE]
    # מנחה של שישי
    yield friday_minha_time
    # פלג המנחה
    yield friday_time[consts.PLAG]
    # בואי כלה
    yield friday_minha_time + consts.COME_SHABAT_DIFF_FROM_MINHA_ACCORDING_TO_PLAG
    # הדלקת נרות
    yield shabat_special_time[consts.SHABAT_ENTER]
    # שקיעה
    yield friday_time[consts.SUN_SET]
    # שיעור בוקר שבת
    yield consts.MORNING_LESSON_TIME
    # שחרית של שבת
    yield Time(consts.MORNING_LESSON_TIME) + consts.MORNING_LESSON_DURATION
    # סוף זמן ק"ש
    yield f'{shabat_time[consts.FIRST_SHMA]} {shabat_time[consts.SECOND_SHMA]}'
    # אבות ובנים
    yield fathers_and_sons_time
    # שיעור אחה"צ שבת
    yield fathers_and_sons_time + consts.FATHERS_AND_SONS_DURATION
    # מנחה של שבת
    yield shabat_time[consts.SUN_SET] - consts.SHABAT_MINHA_DURATION
    # צאת שבת רש"י
    yield shabat_special_time[consts.SHABAT_END]
    # צאת שבת ר"ת
    yield consts.FIELD_TO_FILL


def get_fields_titles_according_to_shkia():
    yield 'פרשת השבוע'
    yield 'תאריך'
    yield 'תאריך לועזי'
    yield 'בואי כלה'
    yield 'מנחה של שישי'
    yield 'הדלקת נרות'
    yield 'שקיעה'
    yield 'צאת הכוכבים'
    yield 'שיעור בוקר שבת'
    yield 'שחרית של שבת'
    yield 'סוף זמן ק"ש'
    yield 'אבות ובנים'
    yield 'שיעור אחה"צ שבת'
    yield 'מנחה של שבת'
    yield 'צאת שבת רש"י (גאונים)'
    yield 'צאת שבת ר"ת'


def get_fields_to_write_according_to_shkia(
        friday_time, shabat_time, shabat_special_time, month, friday_times, shabat_times
):
    come_shabat_time = calculate_come_shabat_time_according_to_shkia(friday_times)
    fathers_and_sons_time = get_fathers_and_sons_time(shabat_times)
    # פרשות
    yield shabat_special_time[consts.PARASHA]
    # תאריך עברי
    yield get_date_string(shabat_time[consts.DAY_IN_MONTH], month)
    # תאריך לועזי
    yield shabat_time[consts.WEIRD_DATE]
    # בואי כלה
    yield come_shabat_time
    # מנחה של שישי
    yield come_shabat_time - consts.COME_SHABAT_DIFF_FROM_MINHA_ACCORDING_TO_SHKIA
    # הדלקת נרות
    yield shabat_special_time[consts.SHABAT_ENTER]
    # שקיעה
    yield friday_time[consts.SUN_SET]
    # צאת הכוכבים
    yield friday_time[consts.STARS_OUT]
    # שיעור בוקר שבת
    yield consts.MORNING_LESSON_TIME
    # שחרית של שבת
    yield Time(consts.MORNING_LESSON_TIME) + consts.MORNING_LESSON_DURATION
    # סוף זמן ק"ש
    yield f'{shabat_time[consts.FIRST_SHMA]} {shabat_time[consts.SECOND_SHMA]}'
    # אבות ובנים
    yield fathers_and_sons_time
    # שיעור אחה"צ שבת
    yield fathers_and_sons_time + consts.FATHERS_AND_SONS_DURATION
    # מנחה של שבת
    yield shabat_time[consts.SUN_SET] - consts.SHABAT_MINHA_DURATION
    # צאת שבת רש"י
    yield shabat_special_time[consts.SHABAT_END]
    # צאת שבת ר"ת
    yield consts.FIELD_TO_FILL


def write_data_to_output(titles_line, fields_lines, moladot_of_month, month, year, mishnaiot_time):
    fields_lines = [i for i in fields_lines]
    if len(fields_lines) == 4:
        document = Document(consts.FOUR_LINES_TEMPLATE)
    else:
        document = Document(consts.FIVE_LINES_TEMPLATE)

    document.paragraphs[2].runs[1].text = month
    document.paragraphs[2].runs[3].text = put_quotes(year)
    document.paragraphs[3].runs[4].text = str(mishnaiot_time)
    document.paragraphs[4].runs[1].text = moladot_of_month
    table_rows = document.tables[0].rows
    for i, title in enumerate(titles_line):
        table_rows[0].cells[i].paragraphs[0].runs[0].text = title
    for i, fields_line in enumerate(fields_lines):
        for j, field in enumerate(fields_line):
            table_rows[i + 1].cells[j].paragraphs[0].runs[0].text = str(field)

    document.save(safe_join_path(consts.TIMES_OUTPUT_FOLDER, f'זמני שבת {month} {year}.docx'))


def main(place=consts.DEFAULT_PLACE, month=consts.DEFAULT_MONTH, year_number=consts.DEFAULT_YEAR):
    place = place or input('what place do you want to save the times of? ').strip()
    month = month or input("what month's times? ").strip()
    year_number = year_number or (consts.THOUSANDS_OF_YEARS_OFFSET + convert_by_gimatria(input("what year's times? ")))
    year = get_year_from_number(year_number - consts.THOUSANDS_OF_YEARS_OFFSET)

    month_number = convert_month_to_month_number(month, year_number)
    next_month_number, next_months_year = get_next_month(month_number, year_number)
    next_month = calculate_month_name(next_month_number, next_months_year)
    place_number = consts.PLACE_DICT[place]

    time_day_list = convert_to_time_of_day(*get_times_as_titles_and_times(
        month_number, place_number, year_number,
    ))
    shabat_times_list = convert_to_time_of_day(*get_shabat_times(place_number, year_number))
    moladot_list = convert_to_time_of_day(*get_start_of_months(year_number))
    shabat_special_times = [one_shabat_times for one_shabat_times in shabat_times_list if is_shabat_in_month(
        one_shabat_times, month
    )]
    month_moladot, = [one_moladot for one_moladot in moladot_list if is_moladot_time_in_month(one_moladot, next_month)]

    shabat_times = get_specific_day_times(consts.SHABAT, time_day_list)
    friday_times = get_specific_day_times(consts.FRIDAY, time_day_list)

    if len(shabat_times) > len(friday_times):
        last_month_number, last_months_year = get_last_month(month_number, year_number)
        last_month_times = convert_to_time_of_day(*get_times_as_titles_and_times(
            last_month_number, place_number, last_months_year
        ))
        last_month_friday_times = get_specific_day_times(consts.FRIDAY, last_month_times)
        friday_times = [last_month_friday_times[-1], *friday_times]

    if not len(shabat_times) == len(shabat_special_times) == len(friday_times):
        raise Exception(
            f'shabat_times_len == {len(shabat_times)}, shabat_special_times_len == {len(shabat_special_times)}, ' +
            f'friday_times_len == {len(friday_times)}'
        )

    if is_according_to_plag(friday_times):
        get_titles = get_fields_titles_according_to_plag
        get_fields = get_fields_to_write_according_to_plag
    else:
        get_titles = get_fields_titles_according_to_shkia
        get_fields = get_fields_to_write_according_to_shkia

    titles_line = get_titles()
    fields_lines = (get_fields(
        friday_times[i], shabat_times[i], shabat_special_times[i], month, friday_times, shabat_times
    ) for i in range(len(shabat_times)))
    moladot_of_month = calculate_moladot_of_month(month_moladot)
    mishnaiot_time = get_fathers_and_sons_time(shabat_times) - consts.MISHNAIOT_DURATION

    write_data_to_output(titles_line, fields_lines, moladot_of_month, month, year, mishnaiot_time)


if __name__ == '__main__':
    main()
